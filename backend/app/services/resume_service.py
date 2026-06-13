# backend/app/services/resume_service.py
#
# Full ATS Resume Optimizer pipeline:
#   1. Extract text from resume PDF
#   2. Clean scraped text from UI garbage/tokens
#   3. Extract keywords from job description (TF-IDF + spaCy)
#   4. Score keyword overlap with structural bonus matrix (ATS score 0–100)
#   5. Use AI to rewrite and optimize the resume (Decoupled Prompt Guard)
#   6. Export to PDF (reportlab) or DOCX (python-docx)
#
import re
import io
import logging
from typing import List, Tuple, AsyncGenerator, Dict, Set

#import spacy
try:
    import spacy
except ImportError:
    spacy = None
from sklearn.feature_extraction.text import TfidfVectorizer
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable
)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.ai_service import ai_service
from app.models.schemas import Message

logger = logging.getLogger(__name__)

# Load spaCy English model safely
# Optional spaCy support
try:
    import spacy

    try:
        nlp = spacy.load("en_core_web_sm")
    except Exception:
        logger.warning("spaCy model unavailable. Continuing without NLP model.")
        nlp = None

except ImportError:
    logger.warning("spaCy not installed. ATS analysis will use TF-IDF only.")
    spacy = None
    nlp = None
``


# Core professional English stop words to wipe from keyphrase generation pipelines
STOP_WORDS_CUSTOM = {
    "who", "which", "part", "skip", "week", "days", "years", "role", "work", 
    "apply", "scroll", "click", "button", "job", "company", "description", 
    "billion", "version", "status", "page", "requirements", "responsibilities"
}

# Skills / tech keywords to always look for
TECH_SKILLS = {
    # Languages
    "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust",
    "swift", "kotlin", "ruby", "php", "scala", "r", "matlab", "html", "css",
    # Frontend
    "react", "vue", "angular", "next.js", "svelte", "tailwind", "bootstrap", "webpack", "vite", "redux",
    # Backend / Engineering Frameworks
    "node.js", "express", "fastapi", "django", "flask", "spring", "graphql", "rest api", "microservices", "grpc", ".net",
    # Databases & Platforms
    "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "supabase", "firebase", "dynamodb", "sqlite", "finacle",
    # Cloud / DevOps
    "aws", "gcp", "azure", "docker", "kubernetes", "ci/cd", "github actions", "terraform", "ansible", "linux", "nginx",
    # AI / ML
    "machine learning", "deep learning", "tensorflow", "pytorch", "langchain", "llm", "openai", "nlp", "hugging face",
    # Methodologies
    "git", "jira", "figma", "agile", "scrum", "testing", "banking"
}


class ResumeService:

    # ─────────────────────────────────────────────────────────────────────────
    # TEXT EXTRACTION & CLEANSING UTILITIES
    # ─────────────────────────────────────────────────────────────────────────

    def extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """Extract all text from a PDF file."""
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        return "\n\n".join(pages)

    def _clean_text(self, raw_text: str) -> str:
        """
        Removes HTML remnants, standalone integers, and scraping garbage strings 
        before sending text data to the tokenization processors.
        """
        if not raw_text:
            return ""
        
        # Strip raw HTML tags and line breaks
        text = re.sub(r'<[^>]+>', ' ', raw_text)
        text = re.sub(r'\b\d+br\b', ' ', text, flags=re.IGNORECASE)
        
        # Remove standalone web portal UI navigation string phrases
        ui_phrases = [
            r'apply to job', r'scroll to top', r'privacy statement', r'save job', 
            r'send to friend', r'create account', r'sign in', r'all rights reserved'
        ]
        for phrase in ui_phrases:
            text = re.sub(phrase, ' ', text, flags=re.IGNORECASE)
            
        # Strip standalone integers / random scraping index codes (e.g. 5179, 26209)
        text = re.sub(r'\b\d{3,7}\b', ' ', text)
        
        return text.strip()

    def parse_contact_information(self, text: str) -> Dict[str, bool]:
        """
        Advanced contact parsing block matching international formats, 
        Indian country codes (+91), email addresses, and LinkedIn profiles.
        """
        text_lower = text.lower()
        
        # Supported layouts: +91 7013128228, +91-70131-28228, 07013128228, 7013128228
        phone_pattern = r'(?:\+?91[\s\-\.]?)?0?\d{10}\b'
        email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        linkedin_pattern = r'linkedin\.com\/in\/[\w\-\_]+'

        has_phone = bool(re.search(phone_pattern, text))
        has_email = bool(re.search(email_pattern, text_lower))
        has_linkedin = bool(re.search(linkedin_pattern, text_lower))

        return {
            "has_phone": has_phone,
            "has_email": has_email,
            "has_linkedin": has_linkedin,
            "is_compliant": has_phone and has_email
        }

    # ─────────────────────────────────────────────────────────────────────────
    # KEYWORD EXTRACTION PIPELINE
    # ─────────────────────────────────────────────────────────────────────────

    def extract_keywords(self, text: str, top_n: int = 40) -> List[str]:
        """
        Extracts verified semantic terms from raw target fields. 
        Applies cleaning filters, noun chunk filters, and stop word scrubs.
        """
        cleaned_text = self._clean_text(text)
        keywords = set()

        if not cleaned_text.strip():
            return []

        # Method 1: TF-IDF on character n-grams catches compound terms
        try:
            vectorizer = TfidfVectorizer(
                ngram_range=(1, 2),
                stop_words="english",
                max_features=150,
                analyzer="word",
            )
            vectorizer.fit([cleaned_text])
            scores = dict(zip(vectorizer.get_feature_names_out(), vectorizer.idf_))
            top_terms = sorted(scores.items(), key=lambda x: x[1])[:top_n]
            
            for term, _ in top_terms:
                term_clean = term.lower().strip()
                if len(term_clean) > 2 and not term_clean.isnumeric() and term_clean not in STOP_WORDS_CUSTOM:
                    keywords.add(term_clean)
        except Exception as e:
            logger.warning(f"TF-IDF extraction skipped: {e}")

        # Method 2: spaCy NLP Proper Noun and Named Entity Extraction Layer
        if nlp:
            try:
                doc = nlp(cleaned_text[:10000])
                for chunk in doc.noun_chunks:
                    chunk_text = chunk.text.lower().strip()
                    # Wipe single chars, pure numeric codes, and custom stop word strings
                    if 2 < len(chunk_text) < 40 and not chunk_text.isnumeric():
                        if not any(stop in chunk_text.split() for stop in STOP_WORDS_CUSTOM):
                            keywords.add(chunk_text)
                            
                for ent in doc.ents:
                    if ent.label_ in ("ORG", "PRODUCT", "LANGUAGE", "GPE"):
                        ent_text = ent.text.lower().strip()
                        if len(ent_text) > 2 and ent_text not in STOP_WORDS_CUSTOM:
                            keywords.add(ent_text)
            except Exception as e:
                logger.warning(f"spaCy engine extraction skipped: {e}")

        # Method 3: Core technology index matching layer validation loop
        text_lower = cleaned_text.lower()
        for skill in TECH_SKILLS:
            if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                keywords.add(skill)

        return sorted(list(keywords))

    # ─────────────────────────────────────────────────────────────────────────
    # MULTI-FACTOR ATS SCORING MATRIX
    # ─────────────────────────────────────────────────────────────────────────

    def calculate_ats_score(
        self,
        resume_text: str,
        jd_keywords: List[str],
    ) -> Tuple[int, List[str], List[str]]:
        """
        Multi-factor compliance evaluation matrix scaling structural presence, 
        contact availability, and clean keyword matching logic.
        """
        if not jd_keywords:
            return 0, [], []

        resume_lower = resume_text.lower()
        matched = []
        missing = []

        # 1. Evaluate clean keyword coverage tracking rows
        for keyword in jd_keywords:
            pattern = r'\b' + re.escape(keyword.lower()) + r'\b'
            if re.search(pattern, resume_lower):
                matched.append(keyword)
            else:
                missing.append(keyword)

        # Base Weight: 50% max allocated to matching structural keywords
        keyword_ratio = len(matched) / len(jd_keywords)
        weighted_base = keyword_ratio * 50

        # 2. Section Compliance Matrix Check: 35% max allocation
        sections = {
            "summary|objective|profile": 5,
            "experience|work history|employment": 10,
            "education|degree|university|college": 10,
            "skills|technologies|tools|proficiencies": 5,
            "projects|portfolio": 5
        }
        section_bonus = sum(
            pts for pattern, pts in sections.items() 
            if re.search(pattern, resume_lower, re.IGNORECASE)
        )

        # 3. Contact Compliance Verification Check: 15% max allocation
        contact_metrics = self.parse_contact_information(resume_text)
        contact_bonus = 0
        if contact_metrics["has_phone"]: contact_bonus += 5
        if contact_metrics["has_email"]: contact_bonus += 5
        if contact_metrics["has_linkedin"]: contact_bonus += 5

        # Compile total score balances
        final_score = int(weighted_base + section_bonus + contact_bonus)
        final_score = max(0, min(100, final_score))

        # Sort missing keywords for UI delivery presentation tracks
        missing.sort(key=len)

        return final_score, matched, missing

    # ─────────────────────────────────────────────────────────────────────────
    # DECOUPLED ANTI-CONTAMINATION PROMPT GUARD REWRITE
    # ─────────────────────────────────────────────────────────────────────────

    async def rewrite_resume(
        self,
        resume_text: str,
        job_description: str,
        missing_keywords: List[str],
        matched_keywords: List[str],
        ats_score: int,
        model: str = "llama3",
    ) -> AsyncGenerator[str, None]:
        """
        Uses an LLM prompt guard block to optimize resume bullet points.
        Prevents job description data from bleeding into history bullet points,
        and cuts all internal meta-commentary or improvement logs.
        """

        missing_str = ", ".join(missing_keywords[:20]) if missing_keywords else "None"
        matched_str = ", ".join(matched_keywords[:15]) if matched_keywords else "None"

        prompt = f"""You are a professional executive resume writer. Your task is to optimize the candidate's historical resume statements against the target requirements.

CRITICAL EXECUTION CONSTRAINTS:
1. SOURCE ISOLATION: Use the Job Description ONLY to understand the target skills, terminology framing, and tech keywords. NEVER copy sentences, responsibilities, or selection criteria text verbatim from the Job Description into the resume.
2. NO HALLUCINATION: Optimize and rephrase the candidate's existing work bullets into strong action statements. Do NOT fabricate entirely new employers, contract timelines, or advanced degrees.
3. CONTEXT INTEGRITY: Do not inventory skills the candidate does not hint at in their base text. Do not add highly specific modern engines unless backed up by their technical history.
4. NO META-COMMENTARY: Output ONLY the final resume. Do NOT include text like "Key Improvements Made", "Changes Applied", explanations of your choices, or cover letter closing blocks at the bottom.

CURRENT ATS RATING: {ats_score}/100
MISSING TARGET SKILLS TO INTEGRATE NATURALLY: {missing_str}
EXISTING MATCHED SKILLS: {matched_str}

TARGET JOB DESCRIPTION:
{job_description[:2000]}

CANDIDATE BASE RESUME TEXT:
{resume_text[:3000]}

Generate the output using these exact markdown headers:

## PROFESSIONAL SUMMARY
[Generate a powerful, 3-sentence summary tailored to this target profile. Lead with the target role name.]

## SKILLS
[Categorized technical skills section including the missing integrated skills.]

## EXPERIENCE
[Rewritten work entries maximizing metric outcomes and action verbs. Retain original timeline dates.]

## EDUCATION
[Formatted clean text layout of the candidate's academic milestones.]

## PROJECTS
[Rewritten projects section incorporating the candidate's technical achievements and stack.]"""

        messages = [Message(role="user", content=prompt)]
        async for token in ai_service.stream_response(messages, model):
            yield token

    async def generate_suggestions(
        self,
        resume_text: str,
        job_description: str,
        ats_score: int,
        model: str = "llama3",
    ) -> AsyncGenerator[str, None]:
        """Generate concise, actionable improvement suggestions."""

        prompt = f"""ATS Score: {ats_score}/100.

Resume:
{resume_text[:1500]}

Job Description:
{job_description[:800]}

List exactly 8 specific, actionable suggestions to improve this resume for ATS and human readers.
Format each as: "ACTION: specific thing to do (reason why it helps ATS/hiring managers)"
Be direct, specific, and practical. No generic advice."""

        messages = [Message(role="user", content=prompt)]
        async for token in ai_service.stream_response(messages, model):
            yield token

    # ─────────────────────────────────────────────────────────────────────────
    # EXPORT: PDF via reportlab
    # ─────────────────────────────────────────────────────────────────────────

    def export_to_pdf(self, resume_content: str, filename: str = "resume") -> bytes:
        """Renders the rewritten resume as a clean, ATS-friendly PDF."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()

        name_style = ParagraphStyle(
            "Name", parent=styles["Normal"], fontSize=20, fontName="Helvetica-Bold", textColor=colors.HexColor("#0f172a"), spaceAfter=4, leading=24
        )
        section_style = ParagraphStyle(
            "Section", parent=styles["Normal"], fontSize=12, fontName="Helvetica-Bold", textColor=colors.HexColor("#1e3a8a"), spaceBefore=14, spaceAfter=4, leading=16
        )
        body_style = ParagraphStyle(
            "Body", parent=styles["Normal"], fontSize=10, fontName="Helvetica", textColor=colors.HexColor("#334155"), spaceAfter=3, leading=14
        )
        bullet_style = ParagraphStyle(
            "Bullet", parent=body_style, leftIndent=16, firstLineIndent=-10, spaceAfter=2, leading=14
        )

        story = []

        cleaned_content = re.sub(r'(?<!#)##([A-Za-z])', r'## \1', resume_content)
        cleaned_content = re.sub(r'([a-z])([A-Z])', r'\1 \2', cleaned_content)
        
        lines = cleaned_content.split("\n")

        for line in lines:
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 4))
                continue

            if stripped.startswith("##") or "PROFESSIONAL SUMMARY" in stripped.upper() or "EXPERIENCE" in stripped.upper() or "SKILLS" in stripped.upper() or "PROJECTS" in stripped.upper():
                header = stripped.replace("#", "").strip()
                story.append(Spacer(1, 6))
                story.append(Paragraph(header.upper(), section_style))
                story.append(HRFlowable(
                    width="100%",
                    thickness=1.0,
                    color=colors.HexColor("#1e3a8a"),
                    spaceAfter=6,
                ))
            elif stripped.startswith("- ") or stripped.startswith("• "):
                text = stripped[2:].strip()
                story.append(Paragraph(f"• {text}", bullet_style))
            elif re.match(r"^[A-Z][A-Z\s]+$", stripped) and len(stripped) < 60:
                story.append(Paragraph(stripped, name_style))
            else:
                story.append(Paragraph(stripped, body_style))

        doc.build(story)
        return buffer.getvalue()

    # ─────────────────────────────────────────────────────────────────────────
    # EXPORT: DOCX via python-docx
    # ─────────────────────────────────────────────────────────────────────────

    def export_to_docx(self, resume_content: str) -> bytes:
        """Generates an ATS-friendly .docx file with standard core styles."""
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        cleaned_content = re.sub(r'(?<!#)##([A-Za-z])', r'## \1', resume_content)
        cleaned_content = re.sub(r'([a-z])([A-Z])', r'\1 \2', cleaned_content)

        lines = cleaned_content.split("\n")

        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue

            if stripped.startswith("##") or "SUMMARY" in stripped.upper() or "EXPERIENCE" in stripped.upper() or "SKILLS" in stripped.upper():
                header = stripped.replace("#", "").strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(header.upper())
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(30, 58, 138)
            elif stripped.startswith("- ") or stripped.startswith("• "):
                text = stripped[2:].strip()
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(text)
                run.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(2)
            elif re.match(r"^[A-Z][A-Z\s]+$", stripped) and len(stripped) < 60:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(20)
                run.font.color.rgb = RGBColor(15, 23, 42)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                p.paragraph_format.space_after = Pt(3)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()


resume_service = ResumeService()